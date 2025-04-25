import streamlit as st
import cv2
import numpy as np
import os
import tempfile
import time
import torch
import whisper
import pandas as pd
from io import BytesIO
from pydub import AudioSegment
from indic_transliteration import sanscript
from indic_transliteration.sanscript import transliterate
from PIL import Image
from deepface import DeepFace
import json
from docx import Document
import base64
from pyannote.audio import Pipeline


# Custom CSS for styling
st.markdown("""
    <style>
        /* Remove scrollbar */
        .main .block-container {
            max-width: 100%;
            padding: 2rem 2rem 2rem;
        }
        /* Hide scrollbar */
        ::-webkit-scrollbar {
            display: none;
        }
        html {
            overflow: hidden;
        }
        body {
            overflow: hidden;
        }
        .main {
            max-width: 500px;
            margin: 0 auto;
            padding: 20px;
        }
        .square-box {
            border: 2px solid #ddd;
            border-radius: 40px;
            padding: 20px;
            text-align: center;
            margin: 10px auto;
            background-color: #f9f9f9;
            cursor: pointer;
        }
        .square-box:hover {
            background-color: #e6e6e6;
        }
        .button {
            padding: 10px 20px;
            margin: 10px;
            font-size: 16px;
        }
        .patient-card {
            border: 1px solid #ddd;
            border-radius: 10px;
            padding: 15px;
            margin-bottom: 15px;
            background-color: #f8f9fa;
        }
        .file-item {
            padding: 8px;
            margin: 5px 0;
            background-color: #f0f0f0;
            border-radius: 5px;
            display: flex;
            justify-content: space-between;
            align-items: center;
        }
        .file-item button {
            margin-left: 10px;
        }
        .audio-player {
            width: 100%;
            margin: 10px 0;
        }
        .patient-name {
            cursor: pointer;
            color: #1e88e5;
            font-weight: bold;
        }
        .patient-name:hover {
            text-decoration: underline;
        }
        .doctor-text {
            color: #1e88e5;
            font-weight: bold;
            background-color: #e6f2ff;
            padding: 8px;
            border-radius: 5px;
            margin: 5px 0;
        }
        .patient-text {
            color: #2e7d32;
            font-weight: bold;
            background-color: #e8f5e9;
            padding: 8px;
            border-radius: 5px;
            margin: 5px 0;
        }
        .other-text {
            color: #6a1b9a;
            font-weight: bold;
            background-color: #f3e5f5;
            padding: 8px;
            border-radius: 5px;
            margin: 5px 0;
        }
    </style>
""", unsafe_allow_html=True)

# ----------------- NAVIGATION FUNCTION -----------------
def navigate(page):
    st.session_state.page = page
    st.rerun() 

# ----------------- LOGIN PAGE -----------------
def login_page():
    st.image("https://glimageurl.golocall.com/golocal-post/image/920535_newliferehabanddeaddictioncentre_1594791897.jpeg")
    st.title("*WELCOME !*")

    name = st.text_input("Enter your name", key="login_name")
    password = st.text_input("Enter password", type="password", key="login_password")

    if st.button("Login"):
        error = False

        # ✅ Name validation
        invalid_names = {"xyz", "abc", "test", "demo", "n/a"}
        if not name.strip():
            st.error("Name is required and cannot be empty.")
            error = True
        elif not name.replace(" ", "").isalpha():
            st.error("Name must contain only alphabets.")
            error = True
        elif name.lower() in invalid_names:
            st.error(f"'{name}' is not a valid name.")
            error = True
        elif len(name.split()) < 2:
            st.error("Please provide a full name (e.g., Firstname Lastname).")
            error = True

        # ✅ Password validation
        special_chars = "!@#$%^&*()_+[]{}|;:,.<>?/~"
        if len(password) < 8 or not any(char in special_chars for char in password):
            st.error("Password must be at least 8 characters long and contain one special character.")
            error = True

        if not error:
            st.session_state.page = 'dashboard'
            st.success("Login successful! ✅")
            st.rerun()

# ----------------- DASHBOARD PAGE -----------------
def dashboard_page():
    st.sidebar.title("Navigation")
    if st.sidebar.button("🏠 Dashboard"):
        navigate("dashboard")
    if st.sidebar.button("🎙️ Record Audio"):
        navigate("record_audio")
    if st.sidebar.button("📝 Transcribe Audio"):
        navigate("transcribe_audio")
    if st.sidebar.button("📤 Upload File"):
        navigate("upload_file")
    if st.sidebar.button("📋 View Details"):
        navigate("view_details")
    if st.sidebar.button("🚪 Logout"):
        navigate("login")

    st.title("Hello!")
    st.write("Every mind has a story. Healing begins when we listen.")

    col1, col2 = st.columns(2)
    with col1:
        st.image("https://easy-peasy.ai/cdn-cgi/image/quality=80,format=auto,width=700/https://fdczvxmwwjwpwbeeqcth.supabase.co/storage/v1/object/public/images/fffd9126-dda4-430c-a18d-fb33c6493c57/de210368-9622-4654-b8c7-a7f24673cb00.png", width=200, caption="Doctor")
        if st.button("Doctor Section"):
            navigate("doctor_page")
    with col2:
        st.image("https://banner2.cleanpng.com/20190422/fe/kisspng-illustration-clip-art-human-behavior-child-teachers39-domain-japanese-joy-5cbdb28aecaaa4.1109688515559358829694.jpg", width=100, caption="Patient")
        if st.button("Patient Section"):
            navigate("patient_page")

# ----------------- DATA PERSISTENCE FUNCTIONS -----------------
def load_data():
    """Load patient and doctor data from JSON files"""
    try:
        with open('patients.json', 'r') as f:
            patients_data = json.load(f)
            # Convert base64 encoded data back to bytes
            for patient in patients_data:
                if 'records_data' in patient and patient['records_data']:
                    patient['records_data'] = base64.b64decode(patient['records_data'].encode('utf-8'))
                if 'audio_files' in patient:
                    for audio_file in patient['audio_files']:
                        if 'file_data' in audio_file and audio_file['file_data']:
                            audio_file['file_data'] = base64.b64decode(audio_file['file_data'].encode('utf-8'))
                if 'transcriptions' in patient:
                    for transcription in patient['transcriptions']:
                        if 'file_data' in transcription and transcription['file_data']:
                            transcription['file_data'] = base64.b64decode(transcription['file_data'].encode('utf-8'))
            st.session_state.patients = patients_data
    except (FileNotFoundError, json.JSONDecodeError):
        st.session_state.patients = []
    
    try:
        with open('doctors.json', 'r') as f:
            st.session_state.doctors = json.load(f)
    except (FileNotFoundError, json.JSONDecodeError):
        st.session_state.doctors = ["John Doe", "Jane Smith"]  # Default doctors

def save_patients():
    """Save patient data to JSON file"""
    patients_to_save = []
    for patient in st.session_state.patients:
        patient_copy = patient.copy()
        # Convert binary data to base64 for JSON serialization
        if 'records_data' in patient_copy:
            patient_copy['records_data'] = base64.b64encode(patient_copy.get('records_data', b'')).decode('utf-8') if patient_copy.get('records_data') else None
        if 'audio_files' in patient_copy:
            for audio_file in patient_copy['audio_files']:
                if 'file_data' in audio_file:
                    audio_file['file_data'] = base64.b64encode(audio_file.get('file_data', b'')).decode('utf-8') if audio_file.get('file_data') else None
        if 'transcriptions' in patient_copy:
            for transcription in patient_copy['transcriptions']:
                if 'file_data' in transcription:
                    transcription['file_data'] = base64.b64encode(transcription.get('file_data', b'')).decode('utf-8') if transcription.get('file_data') else None
        patients_to_save.append(patient_copy)
    
    with open('patients.json', 'w') as f:
        json.dump(patients_to_save, f)

def save_doctors():
    """Save doctor data to JSON file"""
    with open('doctors.json', 'w') as f:
        json.dump(st.session_state.doctors, f)

def get_file_icon(filename):
    """Get appropriate icon based on file extension"""
    if not filename:
        return "📄"
    ext = os.path.splitext(filename)[1].lower()
    if ext in ['.jpg', '.jpeg', '.png', '.gif']:
        return "🖼️"
    elif ext in ['.pdf']:
        return "📄"
    elif ext in ['.doc', '.docx']:
        return "📝"
    elif ext in ['.mp3', '.wav', '.m4a']:
        return "🎵"
    elif ext in ['.txt']:
        return "📋"
    else:
        return "📁"

# ----------------- CAPTURE AND ANALYZE -----------------
def capture_and_analyze():
    st.title("🧠 Depression & Stress Analysis from Facial Expressions")
    st.write("📷 Capture an image using your camera to analyze depression, stress, and emotions.")

    img_file = st.camera_input("Capture Image")

    if img_file is not None:
        try:
            # Load the captured image
            image = Image.open(img_file)
            img_array = np.array(image)  # Convert PIL image to NumPy array
            
            # Convert RGB to BGR for OpenCV
            img_bgr = cv2.cvtColor(img_array, cv2.COLOR_RGB2BGR)

            # Display the captured image
            st.image(image, caption="Captured Image", use_column_width=True)

            # Save a temporary image file for DeepFace
            temp_path = "captured_face.jpg"
            cv2.imwrite(temp_path, img_bgr)  # Save image

            with st.spinner("🧐 Analyzing emotions and stress levels..."):
                # Detect Emotion using DeepFace
                try:
                    result = DeepFace.analyze(img_path=temp_path, actions=['emotion'], enforce_detection=False)
                    if isinstance(result, list) and len(result) > 0:
                        result = result[0]  # Extract first face analysis

                    emotion = result.get('dominant_emotion', 'Unknown')
                    emotion_scores = result.get('emotion', {})

                    # Display emotion results
                    st.success(f"🎭 Detected Emotion: *{emotion.capitalize()}*")
                    st.write("🔍 Emotion Scores:", emotion_scores)

                    # Calculate Depression Level
                    depression_level = round(emotion_scores.get('sad', 0) + emotion_scores.get('fear', 0), 2)
                    st.write(f"😞 *Depression Level:* {depression_level}%")

                except Exception as e:
                    st.error(f"⚠️ Emotion analysis error: {str(e)}")

                # Detect Eye Blinking using OpenCV
                try:
                    face_cascade = cv2.CascadeClassifier(cv2.data.haarcascades + 'haarcascade_frontalface_default.xml')
                    eye_cascade = cv2.CascadeClassifier(cv2.data.haarcascades + 'haarcascade_eye.xml')

                    gray = cv2.cvtColor(img_bgr, cv2.COLOR_BGR2GRAY)
                    faces = face_cascade.detectMultiScale(gray, 1.1, 4)

                    blink_score = 0
                    for (x, y, w, h) in faces:
                        roi_gray = gray[y:y + h, x:x + w]
                        eyes = eye_cascade.detectMultiScale(roi_gray)
                        blink_score = len(eyes)

                    stress_level = min(100, blink_score * 10)  # Convert eye count to stress percentage
                    st.write(f"😰 *Stress Level (Eye Blink Analysis):* {stress_level}%")

                except Exception as e:
                    st.error(f"⚠️ Eye blink detection error: {str(e)}")

        except Exception as e:
            st.error(f"❌ Error: {str(e)}")

    if st.button("Done"):
        navigate("dashboard")

# ----------------- DOCTOR PAGE -----------------
def doctor_page():
    st.title("Doctors List")
    
    # Input for new doctor name
    new_doctor = st.text_input("Add New Doctor (Enter full name)")
    if st.button("Save Doctor"):
        if new_doctor.strip():
            if new_doctor not in st.session_state.doctors:
                st.session_state.doctors.append(new_doctor)
                save_doctors()
                st.success(f"Dr. {new_doctor} added successfully!")
                st.rerun()
            else:
                st.warning("This doctor is already in the list")
        else:
            st.error("Please enter a valid name")
    
    # Display current doctors with delete buttons
    st.write("### Current Doctors")
    
    # Create a form for the delete functionality
    with st.form(key="doctors_form"):
        # Display all doctors with checkboxes for deletion
        doctors_to_delete = []
        for doctor in st.session_state.doctors:
            # Use checkbox for each doctor with the doctor's name as label
            delete = st.checkbox(f"Dr. {doctor}", key=f"delete_{doctor}")
            if delete:
                doctors_to_delete.append(doctor)
        
        # Single submit button for all deletions
        if st.form_submit_button("🗑️"):
            if doctors_to_delete:
                # Remove selected doctors
                st.session_state.doctors = [doc for doc in st.session_state.doctors if doc not in doctors_to_delete]
                save_doctors()
                st.success(f"Deleted {len(doctors_to_delete)} doctor(s)")
                st.rerun()
            else:
                st.warning("Please select at least one doctor to delete")
    
    if st.button("Back"):
        navigate("dashboard")

# ----------------- PATIENT PAGE -----------------
def patient_page():
    st.title("Patients List")
    if st.button("New Patient"):
        navigate("new_patient")
    if st.button("Existing Patients"):
        navigate("existing_patients_page")
    if st.button("Back"):
        navigate("dashboard")

# ----------------- NEW PATIENT PAGE -----------------
def new_patient_page():
    name = st.text_input("Name (Required*)")
    age = st.text_input("Age (Required*)")
    mobile = st.text_input("Mobile Number (Required*)")
    history = st.text_area("Medical History (Optional)")
    symptoms = st.text_area("Symptoms (Optional)")
    records = st.file_uploader("Upload Medical Records (Optional - Any file type allowed)", type=None)

    if st.button("Submit"):
        if name and age.isdigit() and mobile.isdigit() and len(mobile) == 10:
            patient_data = {
                "name": name,
                "age": int(age),
                "mobile": mobile,
                "history": history,
                "symptoms": symptoms,
                "records": records.name if records else None,
                "records_data": records.read() if records else None,
                "audio_files": [],
                "transcriptions": [],
                "analysis_data": []
            }
            st.session_state.patients.append(patient_data)
            save_patients()  # Save to file
            st.success("Patient added successfully! ✅")
            navigate("existing_patients_page")
        else:
            st.error("❌ Invalid input. Please check your entries.")

# ----------------- EXISTING PATIENTS PAGE -----------------
def existing_patients_page():
    st.title("Existing Patients")
    
    if not st.session_state.patients:
        st.warning("No patients found.")
    else:
        for idx, patient in enumerate(st.session_state.patients):
            col1, col2, col3 = st.columns([3, 1, 1])
            with col1:
                st.write(f"{patient['name']} (Age: {patient['age']})")
            with col2:
                if st.button(f"📷", key=f"analyze_{idx}"):
                    st.session_state.selected_patient = patient
                    navigate("image_analysis")
            with col3:
                if st.button(f"🗑️", key=f"delete_{idx}"):
                    # Remove patient from list
                    st.session_state.patients.pop(idx)
                    save_patients()  # Save to file
                    st.rerun()  # Refresh the page to show updated list
    
    if st.button("Back"):
        navigate("dashboard")

# ----------------- VIEW DETAILS PAGE -----------------
def view_details_page():
    st.title("Patient Details")
    
    if not st.session_state.patients:
        st.warning("No patient records found.")
    else:
        # Display clickable patient names
        st.subheader("Select a Patient")
        for patient in st.session_state.patients:
            if st.button(f"👤 {patient['name']} (Age: {patient['age']})", key=f"patient_{patient['name']}"):
                st.session_state.selected_patient = patient
        
        # Display details of selected patient
        if 'selected_patient' in st.session_state and st.session_state.selected_patient:
            patient = st.session_state.selected_patient
            with st.container():
                st.markdown(f"""
                <div class="patient-card">
                    <h3>{patient['name']}</h3>
                    <p><strong>Age:</strong> {patient['age']}</p>
                    <p><strong>Mobile:</strong> {patient['mobile']}</p>
                    <p><strong>Medical History:</strong> {patient['history'] or 'Not provided'}</p>
                    <p><strong>Symptoms:</strong> {patient['symptoms'] or 'Not provided'}</p>
                </div>
                """, unsafe_allow_html=True)
                
                # Show uploaded medical records if available
                if patient.get('records'):
                    st.subheader("📁 Medical Records")
                    if patient.get('records_data'):
                        col1, col2 = st.columns([3, 1])
                        with col1:
                            st.write(f"{get_file_icon(patient['records'])} {patient['records']}")
                        with col2:
                            st.download_button(
                                label="Download",
                                data=patient['records_data'],
                                file_name=patient['records'],
                                mime="application/octet-stream",
                                key=f"records_{patient['name']}"
                            )
                
                # Show audio files if available
                if patient.get('audio_files'):
                    st.subheader("🎤 Audio Recordings")
                    for idx, audio_file in enumerate(patient['audio_files']):
                        with st.expander(f"{get_file_icon(audio_file['filename'])} {audio_file['filename']} ({audio_file['date']})"):
                            if audio_file.get('file_data'):
                                st.audio(audio_file['file_data'], format='audio/wav')
                                st.download_button(
                                    label="Download Audio",
                                    data=audio_file['file_data'],
                                    file_name=audio_file['filename'],
                                    mime="audio/wav",
                                    key=f"audio_{patient['name']}_{idx}"
                                )
                
                # Show transcriptions if available
                if patient.get('transcriptions'):
                    st.subheader("📝 Transcriptions")
                    for idx, transcription in enumerate(patient['transcriptions']):
                        with st.expander(f"{get_file_icon(transcription['filename'])} {transcription['filename']} ({transcription['date']})"):
                            if transcription.get('file_data'):
                                ext = os.path.splitext(transcription['filename'])[1].lower()
                                if ext == '.txt':
                                    st.text(transcription['file_data'].decode('utf-8'))
                                elif ext == '.docx':
                                    st.warning("DOCX content preview not available. Please download to view.")
                                
                                st.download_button(
                                    label="Download Transcription",
                                    data=transcription['file_data'],
                                    file_name=transcription['filename'],
                                    mime="text/plain" if ext == '.txt' else "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                    key=f"trans_{patient['name']}_{idx}"
                                )
                
                # Show analysis data if available
                if patient.get('analysis_data'):
                    st.subheader("Analysis Results")
                    for analysis in patient['analysis_data']:
                        st.write(f"📅 {analysis['date']}")
                        st.write(f"😊 Emotion: {analysis['emotion']}")
                        st.write(f"😞 Depression Level: {analysis['depression_level']}%")
                        st.write(f"😰 Stress Level: {analysis['stress_level']}%")
                        st.write("---")
    
    if st.button("Back to Dashboard"):
        navigate("dashboard")

# ----------------- UPLOAD FILE PAGE -----------------
def upload_file_page():
    st.title("📤 Upload Audio & Transcription File")
    
    if not st.session_state.patients:
        st.warning("No patients available. Please add a patient first.")
        if st.button("Back"):
            navigate("dashboard")
        return
    
    # Patient selection dropdown
    patient_names = [p['name'] for p in st.session_state.patients]
    selected_patient = st.selectbox("Select Patient", patient_names)
    
    audio_file = st.file_uploader("Upload Audio File", type=["mp3", "wav", "m4a"])
    transcription_file = st.file_uploader("Upload Transcription File (Optional)", type=["txt", "docx"])
    
    if st.button("Upload Files"):
        if selected_patient and audio_file:
            # Find the patient and update their records
            for patient in st.session_state.patients:
                if patient['name'] == selected_patient:
                    # Store the audio file data
                    audio_data = {
                        'filename': audio_file.name,
                        'date': time.strftime("%Y-%m-%d %H:%M:%S"),
                        'file_data': audio_file.read()
                    }
                    
                    if 'audio_files' not in patient:
                        patient['audio_files'] = []
                    patient['audio_files'].append(audio_data)
                    
                    # Store transcription if provided
                    if transcription_file:
                        transcription_data = {
                            'filename': transcription_file.name,
                            'date': time.strftime("%Y-%m-%d %H:%M:%S"),
                            'file_data': transcription_file.read()
                        }
                        if 'transcriptions' not in patient:
                            patient['transcriptions'] = []
                        patient['transcriptions'].append(transcription_data)
                    
                    save_patients()  # Save to file
                    st.success("Files Uploaded Successfully!")
                    break
          
        else:
            st.error("Please select a patient and provide an audio file")
    
    if st.button("Back"):
        navigate("dashboard")

# ----------------- TRANSCRIBE AUDIO PAGE -----------------
def transcribe_audio_page():
    st.title("🌍 Multilingual to Structured English Transcription")

    if "uploaded_file" not in st.session_state:
        st.session_state["uploaded_file"] = None
    if "conversation" not in st.session_state:
        st.session_state["conversation"] = pd.DataFrame()

    @st.cache_resource
    def load_models():
        try:
            device = "cuda" if torch.cuda.is_available() else "cpu"
            
            # Load Whisper model (medium for balance of speed/accuracy)
            whisper_model = whisper.load_model("medium", device=device)
            
            # Load speaker diarization model (no-auth version)
            try:
                from speechbrain.pretrained import SpeakerRecognition
                diarization_model = SpeakerRecognition.from_hparams(
                    source="speechbrain/spkrec-ecapa-voxceleb",
                    savedir="tmp_speaker_model"
                )
                return whisper_model, diarization_model
            except:
                return whisper_model, None

        except Exception as e:
            st.error(f"Model loading failed: {str(e)}")
            return None, None

    def process_audio():
        uploaded_file = st.session_state.get("uploaded_file")
        if not uploaded_file:
            return

        with tempfile.NamedTemporaryFile(delete=False, suffix=".wav") as tmp_file:
            tmp_file.write(uploaded_file.getvalue())
            audio_path = tmp_file.name

        st.info("Processing audio... This may take 2-5 minutes...")

        try:
            whisper_model, diarization_model = load_models()
            if not whisper_model:
                raise Exception("Whisper model not loaded")

            # Step 1: Transcribe with Whisper (translate to English)
            result = whisper_model.transcribe(
                audio_path,
                task="translate",  # Convert to English
                word_timestamps=True  # Needed for speaker alignment
            )

            # Step 2: Speaker identification (simplified approach)
            conversation = []
            if diarization_model:
                # If we have speaker model, use it
                try:
                    # Get speaker embeddings (simplified diarization)
                    embeddings = diarization_model.encode_batch(audio_path)
                    # For demo, we'll assume first speaker is doctor, second is patient
                    speakers = ["Doctor", "Patient", "Third Person"]
                    current_speaker = 0
                    
                    for segment in result['segments']:
                        # Alternate speakers based on segment length
                        if len(segment['text']) > 30:  # If long segment, switch speaker
                            current_speaker = (current_speaker + 1) % len(speakers)
                        
                        conversation.append({
                            "Speaker": speakers[current_speaker],
                            "Text": segment['text'],
                            "Start": segment['start'],
                            "End": segment['end']
                        })
                except:
                    # Fallback if speaker diarization fails
                    pass
            
            if not conversation:  # If no diarization, use simple alternation
                speakers = ["Doctor", "Patient"]
                for i, segment in enumerate(result['segments']):
                    conversation.append({
                        "Speaker": speakers[i % 2],
                        "Text": segment['text'],
                        "Start": segment['start'],
                        "End": segment['end']
                    })

            # Create dataframe
            df = pd.DataFrame(conversation)
            st.session_state["conversation"] = df
            st.success("✅ Transcription complete!")

        except Exception as e:
            st.error(f"Processing failed: {str(e)}")
        finally:
            if os.path.exists(audio_path):
                os.remove(audio_path)

    # UI Components
    uploaded_file = st.file_uploader(
        "📤 Upload consultation audio (any language)", 
        type=["wav", "mp3", "m4a", "ogg"]
    )
    
    if uploaded_file:
        st.session_state["uploaded_file"] = uploaded_file
        st.audio(uploaded_file)
        
        if st.button("🚀 Structure Conversation"):
            process_audio()

    if not st.session_state["conversation"].empty:
        st.write("### 🗣 Structured Conversation (English)")
        
        # Display with colored speaker labels
        for _, row in st.session_state["conversation"].iterrows():
            if row["Speaker"] == "Doctor":
                st.markdown(f"""
                <div style='background-color:#e6f7ff; border-left:4px solid #1890ff; 
                    padding:10px; border-radius:0 8px 8px 0; margin:8px 0;'>
                    <b>👨‍⚕ DOCTOR:</b> {row['Text']}<br>
                    <small>{row['Start']:.1f}s - {row['End']:.1f}s</small>
                </div>
                """, unsafe_allow_html=True)
                
            elif row["Speaker"] == "Patient":
                st.markdown(f"""
                <div style='background-color:#f6ffed; border-left:4px solid #52c41a;
                    padding:10px; border-radius:0 8px 8px 0; margin:8px 0;'>
                    <b>👩 PATIENT:</b> {row['Text']}<br>
                    <small>{row['Start']:.1f}s - {row['End']:.1f}s</small>
                </div>
                """, unsafe_allow_html=True)
                
            else:
                st.markdown(f"""
                <div style='background-color:#fff2e8; border-left:4px solid #fa8c16;
                    padding:10px; border-radius:0 8px 8px 0; margin:8px 0;'>
                    <b>👤 {row['Speaker']}:</b> {row['Text']}<br>
                    <small>{row['Start']:.1f}s - {row['End']:.1f}s</small>
                </div>
                """, unsafe_allow_html=True)

        # Export options
        st.write("### 📤 Export Options")
        col1, col2, col3 = st.columns(3)
        
        with col1:
            st.download_button(
                "💾 CSV",
                data=st.session_state["conversation"].to_csv(index=False),
                file_name="consultation.csv",
                mime="text/csv"
            )
        
        with col2:
            doc = Document()
            for _, row in st.session_state["conversation"].iterrows():
                p = doc.add_paragraph(style='Heading2' if row['Speaker'] == 'Doctor' else 'Heading3')
                p.add_run(f"{row['Speaker']}:").bold = True
                p.add_run(f" {row['Text']}")
                doc.add_paragraph(f"({row['Start']:.1f}s - {row['End']:.1f}s)", style='Caption')
            
            bio = BytesIO()
            doc.save(bio)
            st.download_button(
                "📝 Word",
                data=bio.getvalue(),
                file_name="consultation.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        
        with col3:
            html_content = """
            <html><body style="font-family:Arial">
                <h2>Medical Consultation Transcript</h2>
            """
            for _, row in st.session_state["conversation"].iterrows():
                color = "#1890ff" if row['Speaker'] == 'Doctor' else "#52c41a" if row['Speaker'] == 'Patient' else "#fa8c16"
                html_content += f"""
                <div style="margin-bottom:15px;border-left:3px solid {color};padding-left:10px">
                    <b style="color:{color}">{row['Speaker']}:</b>
                    <p>{row['Text']}</p>
                    <small>{row['Start']:.1f}s - {row['End']:.1f}s</small>
                </div>
                """
            html_content += "</body></html>"
            
            st.download_button(
                "🌐 HTML",
                data=html_content,
                file_name="consultation.html",
                mime="text/html"
            )

    if st.button("← Back to Dashboard"):
        navigate("dashboard")
# ----------------- RECORD AUDIO PAGE -----------------            
def record_audio_page():
    st.title("🎙️ Sound Recorder")
    st.write("Click *Start Recording* to capture audio using your microphone.")

    audio_recorder_html = """
        <script>
            let mediaRecorder;
            let audioChunks = [];
            let isRecording = false;
            let isPaused = false;

            function startRecording() {
                navigator.mediaDevices.getUserMedia({ audio: true })
                    .then(stream => {
                        mediaRecorder = new MediaRecorder(stream);
                        mediaRecorder.start();
                        isRecording = true;
                        isPaused = false;
                        audioChunks = [];

                        mediaRecorder.ondataavailable = event => {
                            audioChunks.push(event.data);
                        };

                        mediaRecorder.onstop = () => {
                            let audioBlob = new Blob(audioChunks, { type: 'audio/wav' });
                            let audioUrl = URL.createObjectURL(audioBlob);
                            let audio = new Audio(audioUrl);
                            let downloadLink = document.getElementById('downloadLink');
                            
                            downloadLink.href = audioUrl;
                            downloadLink.download = 'recording.wav';
                            downloadLink.style.display = 'block';
                            
                            let audioPlayer = document.getElementById('audioPlayer');
                            audioPlayer.src = audioUrl;
                            audioPlayer.style.display = 'block';
                        };
                    });
            }

            function pauseRecording() {
                if (isRecording && mediaRecorder.state === "recording") {
                    mediaRecorder.pause();
                    isPaused = true;
                }
            }

            function resumeRecording() {
                if (isRecording && isPaused) {
                    mediaRecorder.resume();
                    isPaused = false;
                }
            }

            function stopRecording() {
                if (isRecording) {
                    mediaRecorder.stop();
                    isRecording = false;
                    isPaused = false;
                }
            }
        </script>

        <div style="margin-bottom: 20px;">
            <button onclick="startRecording()" style="padding: 10px 15px; margin-right: 10px; background-color: #4CAF50; color: white; border: none; border-radius: 4px; cursor: pointer;">🎙️ Start Recording</button>
            <button onclick="pauseRecording()" style="padding: 10px 15px; margin-right: 10px; background-color: #FF9800; color: white; border: none; border-radius: 4px; cursor: pointer;">⏸️ Pause</button>
            <button onclick="resumeRecording()" style="padding: 10px 15px; margin-right: 10px; background-color: #2196F3; color: white; border: none; border-radius: 4px; cursor: pointer;">▶️ Resume</button>
            <button onclick="stopRecording()" style="padding: 10px 15px; background-color: #F44336; color: white; border: none; border-radius: 4px; cursor: pointer;">⏹️ Stop</button>
        </div>
        
        <br><br>
        <audio id="audioPlayer" controls style="display: none;"></audio>
        <a id="downloadLink" style="display: none;">⬇️ Download Recording</a>
    """

    # Render HTML in Streamlit
    st.components.v1.html(audio_recorder_html, height=300)
            
    if st.button("Back to Dashboard"):
        navigate("dashboard")            

# ----------------- MAIN FUNCTION -----------------
def main():
    if 'page' not in st.session_state:
        st.session_state.page = 'login'
    if 'selected_patient' not in st.session_state:
        st.session_state.selected_patient = None
    
    # Load persistent data
    load_data()
    
    # Page routing
    if st.session_state.page == 'login':
        login_page()
    elif st.session_state.page == 'dashboard':
        dashboard_page()
    elif st.session_state.page == 'patient_page':
        patient_page()
    elif st.session_state.page == 'new_patient':
        new_patient_page()
    elif st.session_state.page == 'existing_patients_page':
        existing_patients_page()
    elif st.session_state.page == 'doctor_page':
        doctor_page()
    elif st.session_state.page == 'image_analysis':
        capture_and_analyze()
    elif st.session_state.page == 'record_audio':
        record_audio_page()
    elif st.session_state.page == 'transcribe_audio':
        transcribe_audio_page()
    elif st.session_state.page == 'upload_file':
        upload_file_page()
    elif st.session_state.page == 'view_details':
        view_details_page()

# ----------------- RUN APP -----------------
if __name__ == "__main__":
    main()